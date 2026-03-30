from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Omrani & Taub Call Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Closer Playbook - Get Them to Free Trial')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Call Objective
doc.add_heading('CALL OBJECTIVE', level=1)
doc.add_paragraph('Get Alex Omrani or Michael Taub to say YES to a 7-day free trial of the AI operations team.')

doc.add_paragraph()

# Firm Snapshot
doc.add_heading('FIRM SNAPSHOT (Know This Cold)', level=1)

table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
data = [
    ('Firm', 'Omrani & Taub, P.C.'),
    ('Partners', 'Alex A. Omrani, Michael A. Taub, Isaac A. Arasteh'),
    ('Key Staff', 'Mergime (paralegal - praised in 10+ reviews), Edwin, Avi'),
    ('Locations', 'Manhattan, Queens, Westchester'),
    ('Experience', '25+ years'),
    ('Top Verdict', '$10.58M (child struck by van)'),
    ('Specialty', 'Construction accidents (Labor Law 240/241), vehicle accidents, immigrant workers'),
    ('Phone', '212-714-1515'),
]
for i, (key, val) in enumerate(data):
    table.rows[i].cells[0].text = key
    table.rows[i].cells[1].text = val

doc.add_paragraph()
doc.add_heading('Notable Verdicts to Reference', level=2)
verdicts = [
    '$10.58M - Pedestrian (child struck by van)',
    '$6M - Bus accident',
    '$5M - Construction (demolition laborer)',
    '$4.5M - Construction (undocumented worker - fatal)',
    '$3.9M - Pedestrian (falling stop sign)',
]
for v in verdicts:
    doc.add_paragraph(v, style='List Bullet')

doc.add_paragraph()

# Opening Hook
doc.add_heading('OPENING HOOK (First 60 Seconds)', level=1)
doc.add_paragraph('Do not start with product. Start with their problem.')

p = doc.add_paragraph()
run = p.add_run('Quick question before we dive in - how many settlements did you close last year... and how many of those clients actually left you a Google review?')
run.italic = True

doc.add_paragraph('Wait for answer. They will likely say not many or we do not track that.')

p = doc.add_paragraph()
run = p.add_run('That is exactly what we keep hearing from PI firms doing $5M+ in verdicts. You are winning cases, but Google does not know it. Meanwhile, competitors with half your track record have twice your reviews - and they are getting the calls.')
run.italic = True

doc.add_paragraph()

# Pain Points
doc.add_heading('THE 3 PAIN POINTS TO HIT', level=1)

doc.add_heading('1. Google Reviews Are Being Left on the Table', level=2)
doc.add_paragraph('The Problem:')
problems1 = [
    'They have $10M+ verdicts but their Google presence does not reflect it',
    'Satisfied clients (who praise Mergime, Edwin, Avi in reviews) are never systematically asked',
    'Competitors with weaker track records have more reviews and get more calls',
]
for prob in problems1:
    doc.add_paragraph(prob, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('What to Say: ')
run.bold = True
run2 = p.add_run('Every settlement that closes without a review request is a missed 5-star review, a missed referral, and a missed intake from someone Googling construction accident lawyer NYC tonight.')
run2.italic = True

doc.add_paragraph()
doc.add_heading('2. High-Value Cases Lost After Hours', level=2)
doc.add_paragraph('The Problem:')
problems2 = [
    'Construction accidents happen at 6 AM',
    'Car accidents happen at 11 PM',
    'When a victim searches construction accident lawyer NYC at 2 AM, they call whoever answers first',
    'With $5M+ cases on the line, every missed call is catastrophic',
]
for prob in problems2:
    doc.add_paragraph(prob, style='List Bullet')

doc.add_paragraph()
doc.add_heading('3. Inconsistent Case Qualification', level=2)
doc.add_paragraph('The Problem:')
problems3 = [
    'Labor Law 240 case = $2M+',
    'General negligence case = $200K',
    'The difference depends on asking the right questions: scaffold? ladder? falling object? height?',
    'Inconsistent intake means high-value cases get undervalued or missed',
]
for prob in problems3:
    doc.add_paragraph(prob, style='List Bullet')

doc.add_paragraph()

# Objection Handling
doc.add_heading('OBJECTION HANDLING', level=1)

p = doc.add_paragraph()
run = p.add_run('We are too busy to implement something new')
run.bold = True
doc.add_paragraph('Response: That is exactly why this works - it runs in the background. Mergime and your team do not have to change anything. The AI handles intake, review requests, and client updates automatically.')

p = doc.add_paragraph()
run = p.add_run('We already have a system for reviews')
run.bold = True
doc.add_paragraph('Response: How many reviews did you get last month? And how many settlements closed? That gap is what we fix.')

p = doc.add_paragraph()
run = p.add_run('How do we know it works?')
run.bold = True
doc.add_paragraph('Response: That is why we do a 7-day free trial. You will see real results before you pay anything. If it does not work, you have lost nothing.')

p = doc.add_paragraph()
run = p.add_run('We need to think about it')
run.bold = True
doc.add_paragraph('Response: Every day you wait is another settlement closing without a review request, another 2 AM call going to a competitor. The trial is free. What is the downside?')

doc.add_paragraph()

# The Close
doc.add_heading('CLOSE: THE FREE TRIAL ASK', level=1)
doc.add_paragraph('Do not ask are you interested - assume the next step.')

p = doc.add_paragraph()
run = p.add_run('Here is what I suggest: let us get you set up for a 7-day trial. No cost, no commitment. We will connect to your intake flow, start capturing after-hours inquiries, and trigger review requests on your next few settlements. You will see exactly how it works with your real cases. If it does not deliver, you walk away. Fair?')
run.italic = True

doc.add_paragraph()
doc.add_paragraph('If they hesitate:')
p = doc.add_paragraph()
run = p.add_run('Look - you have built a firm that wins $10M verdicts. But right now, Google does not know that. Every week you wait is another week where competitors with half your track record are getting more calls. Let us just test it. 7 days. What do you have to lose?')
run.italic = True

doc.add_paragraph()

# Numbers
doc.add_heading('NUMBERS TO DROP', level=1)
table2 = doc.add_table(rows=7, cols=2)
table2.style = 'Table Grid'
numbers = [
    ('Metric', 'Value'),
    ('Top verdict', '$10.58M'),
    ('Years experience', '25+'),
    ('Review increase (typical)', '3x in 90 days'),
    ('After-hours inquiry capture', '35%+ of leads'),
    ('Case qualification accuracy', '100% Labor Law 240 identification'),
    ('Status call reduction', '70% fewer calls'),
]
for i, (key, val) in enumerate(numbers):
    table2.rows[i].cells[0].text = key
    table2.rows[i].cells[1].text = val

doc.add_paragraph()

# Before You Hang Up
doc.add_heading('BEFORE YOU HANG UP', level=1)
checklist = [
    'Confirm the trial start date - We will have you live by [DATE]',
    'Identify the point of contact - Who should we coordinate with? Mergime?',
    'Set the follow-up - I will check in on Day 3 to see how the first inquiries are going',
    'Send the recap email - Within 1 hour of the call',
]
for i, item in enumerate(checklist, 1):
    doc.add_paragraph(str(i) + '. ' + item)

doc.add_paragraph()

# Do Not Do This
doc.add_heading('DO NOT DO THIS', level=1)
donts = [
    'Do not trash their current process - they are proud of their verdicts',
    'Do not oversell - let the trial prove it',
    'Do not get technical - they do not care how the AI works, just that it works',
    'Do not let them think about it without a specific follow-up date',
    'Do not forget to mention Mergime by name - she is clearly important to them',
]
for d in donts:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph()

# Final Tip
doc.add_heading('FINAL TIP', level=1)
doc.add_paragraph('This firm wins big cases. They know they are good. The angle is not you need help - it is you are leaving money on the table that you have already earned.')
doc.add_paragraph('Every $10M verdict that does not turn into 5 Google reviews is a missed opportunity. Every 2 AM call that goes to voicemail is a $5M case walking to a competitor.')
p = doc.add_paragraph()
run = p.add_run('The free trial is the easy yes. Get them there.')
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph('Good luck on the call. Close it!')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'c:\Users\Admin\Downloads\OneMind 1\sales-platform\docs\Omrani-Taub-Call-Guide.docx')
print('Word document created successfully!')
